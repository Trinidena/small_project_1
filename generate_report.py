import os
import warnings
import numpy as np
import matplotlib.pyplot as plt
from sklearn.datasets import load_wine
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.svm import SVC
from sklearn.linear_model import LogisticRegression
from matplotlib.colors import ListedColormap
from sklearn.metrics import accuracy_score
from docx import Document
from docx.shared import Inches

# Suppress all warnings
warnings.filterwarnings("ignore")

# Step 1: Get the directory of the current script and create the "report" folder

# Get the directory of the current script
current_dir = os.path.dirname(os.path.abspath(__file__))

# Define the relative folder path within the project directory
project_folder = os.path.join(current_dir, 'report')

# Create the folder if it doesn't exist
if not os.path.exists(project_folder):
    print(f"Creating directory: {project_folder}")
    os.makedirs(project_folder)
else:
    print(f"Directory {project_folder} already exists")


# Step 2: Define the function to plot and save decision regions
def plot_and_save_decision_regions(X, y, classifier, title, filename, resolution=0.02):
    markers = ('s', 'x')
    colors = ('red', 'blue', 'green')
    cmap = ListedColormap(colors[:len(np.unique(y))])

    # Plot decision surface
    x1_min, x1_max = X[:, 0].min() - 1, X[:, 0].max() + 1
    x2_min, x2_max = X[:, 1].min() - 1, X[:, 1].max() + 1
    xx1, xx2 = np.meshgrid(np.arange(x1_min, x1_max, resolution),
                           np.arange(x2_min, x2_max, resolution))
    Z = classifier.predict(np.array([xx1.ravel(), xx2.ravel()]).T)
    Z = Z.reshape(xx1.shape)

    plt.contourf(xx1, xx2, Z, alpha=0.3, cmap=cmap)
    plt.xlim(xx1.min(), xx1.max())
    plt.ylim(xx2.min(), xx2.max())

    # Plot class samples
    for idx, cl in enumerate(np.unique(y)):
        plt.scatter(x=X[y == cl, 0], y=X[y == cl, 1],
                    alpha=0.8, c=colors[idx],
                    marker=markers[idx % 2], label=f'Class {cl}', edgecolor='black')

    plt.xlabel('Alcohol')
    plt.ylabel('Malic Acid')
    plt.title(title)
    plt.legend(loc='upper left')

    # Full path for saving the file
    filepath = os.path.join(project_folder, filename)
    
    # Print the filepath for debugging
    print(f"Saving figure to: {filepath}")
    
    plt.savefig(filepath)
    plt.close()  # Close the plot to avoid displaying in notebook


# Step 3: Define the method to generate bar plots for accuracies
def generate_bar_plots(values, accuracies, model_name, exclude_value=None):
    """
    Generate two bar plots:
    1. With all provided values and accuracies.
    2. Excluding a specific value (if provided).
    
    Args:
    - values: List of hyperparameter values (C or gamma).
    - accuracies: Corresponding accuracy scores for the values.
    - model_name: Name of the model (Logistic Regression or SVM).
    - exclude_value: Value to be excluded in the second plot (optional).
    """
    
    # Plot with all values
    plt.figure(figsize=(8, 6))
    plt.bar([str(v) for v in values], accuracies, color='blue')
    plt.xlabel(f'{model_name} Parameter Values')
    plt.ylabel('Accuracy Score')
    plt.title(f'{model_name} Accuracy for All Values')
    all_values_chart = os.path.join(project_folder, f'{model_name}_accuracy_all.png')
    plt.savefig(all_values_chart)
    plt.close()

    # Plot excluding a specific value if provided
    if exclude_value is not None:
        # Exclude the specified value
        filtered_values = [v for v in values if v != exclude_value]
        filtered_accuracies = [accuracies[i] for i in range(len(values)) if values[i] != exclude_value]
        
        plt.figure(figsize=(8, 6))
        plt.bar([str(v) for v in filtered_values], filtered_accuracies, color='green')
        plt.xlabel(f'{model_name} Parameter Values (Excluding {exclude_value})')
        plt.ylabel('Accuracy Score')
        plt.title(f'{model_name} Accuracy Excluding {exclude_value}')
        filtered_chart = os.path.join(project_folder, f'{model_name}_accuracy_filtered.png')
        plt.savefig(filtered_chart)
        plt.close()

    return all_values_chart, filtered_chart if exclude_value is not None else all_values_chart


# Step 8: Define the method to generate line charts for accuracies
def generate_line_chart(values, accuracies, model_name):
    """
    Generate a line chart for accuracy scores.
    
    Args:
    - values: List of hyperparameter values (C or gamma).
    - accuracies: Corresponding accuracy scores for the values.
    - model_name: Name of the model (Logistic Regression or SVM).
    """
    plt.figure(figsize=(8, 6))
    plt.plot(values, accuracies, marker='o', linestyle='-', color='blue')
    plt.xlabel(f'{model_name} Parameter Values')
    plt.ylabel('Accuracy Score')
    plt.xscale('log')  # Logarithmic scale if you want
    plt.title(f'{model_name} Accuracy Line Chart')
    plt.xlim([10**(-4), 10**4])
    line_chart = os.path.join(project_folder, f'{model_name}_accuracy_line.png')
    plt.savefig(line_chart)
    plt.close()
    
    return line_chart


# Step 4: Load and prepare the Wine dataset

# Load Wine dataset
wine = load_wine()
X = wine.data[:, [0, 1]]  # Selecting only 'Alcohol' and 'Malic Acid' features
y = wine.target

# Standardize the features
scaler = StandardScaler()
X = scaler.fit_transform(X)

# Split the dataset into training and test sets
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)


# Step 5: Logistic Regression Combinations, Plotting, and Accuracy Calculation

logreg_C_values = [0.001, 0.01, 0.1, 1, 10, 100, 1000, 5000]
logreg_accuracies = []

for C in logreg_C_values:
    model = LogisticRegression(C=C, multi_class='ovr', random_state=42)
    model.fit(X_train, y_train)
    
    # Accuracy calculation
    y_pred = model.predict(X_test)
    acc = accuracy_score(y_test, y_pred)
    logreg_accuracies.append(acc)

    # Plot decision regions
    title = f'Logistic Regression with C={C}'
    filename = f'logreg_C_{C}.png'
    plot_and_save_decision_regions(X_train, y_train, classifier=model, title=title, filename=filename)

# Generate bar plots for Logistic Regression
logreg_all_chart, _ = generate_bar_plots(logreg_C_values, logreg_accuracies, 'Logistic Regression')

# Generate line chart for Logistic Regression
logreg_line_chart = generate_line_chart(logreg_C_values, logreg_accuracies, 'Logistic Regression')


# Step 6: SVM Combinations, Plotting, and Accuracy Calculation

svm_gamma_values = [0.001, 0.01, 0.1, 1, 10, 25, 100, 1000]
svm_accuracies = []

for gamma in svm_gamma_values:
    model = SVC(kernel='rbf', gamma=gamma, random_state=42)
    model.fit(X_train, y_train)
    
    # Accuracy calculation
    y_pred = model.predict(X_test)
    acc = accuracy_score(y_test, y_pred)
    svm_accuracies.append(acc)

    # Plot decision regions
    title = f'SVM with RBF Kernel, gamma={gamma}'
    filename = f'svm_gamma_{gamma}.png'
    plot_and_save_decision_regions(X_train, y_train, classifier=model, title=title, filename=filename)

# Generate bar plots for SVM with gamma=100 excluded
svm_all_chart, svm_filtered_chart = generate_bar_plots(svm_gamma_values, svm_accuracies, 'SVM', exclude_value=1000)

# Generate line chart for SVM
svm_line_chart = generate_line_chart(svm_gamma_values, svm_accuracies, 'SVM')


# Step 7: Generate the Report in Word Document

# Create a new Word document
doc = Document()

# Add a title
doc.add_heading('Predictor Comparison Report', 0)

# Add Logistic Regression section with accuracy chart
doc.add_heading('Logistic Regression', level=1)
doc.add_paragraph('Logistic Regression was tested with various regularization strengths (C values). '
                  'As the C value increases, the model tends to overfit, creating tighter decision boundaries. '
                  'For smaller C values, the model becomes more regularized and generalizes better.')

# Insert Logistic Regression accuracy chart
doc.add_heading('Logistic Regression Accuracy', level=2)
doc.add_paragraph('The chart below shows the accuracy for all C values in Logistic Regression.')
doc.add_picture(logreg_all_chart, width=Inches(5.0))

# Insert Logistic Regression line chart
doc.add_heading('Logistic Regression Accuracy Line Chart', level=2)
doc.add_paragraph('The line chart below shows the accuracy across different C values in Logistic Regression.')
doc.add_picture(logreg_line_chart, width=Inches(5.0))

# Insert Logistic Regression decision boundary figures with comments
for C in logreg_C_values:
    doc.add_heading(f'Logistic Regression with C={C}', level=2)
    doc.add_paragraph(f'The decision boundary for Logistic Regression with C={C} is shown below.')
    
    # Path to the figure
    figure_path = os.path.join(project_folder, f'logreg_C_{C}.png')
    doc.add_picture(figure_path, width=Inches(5.0))


# Add SVM section with accuracy chart
doc.add_heading('SVM with RBF Kernel', level=1)
doc.add_paragraph('Support Vector Machines (SVM) with an RBF kernel were tested with different gamma values. '
                  'The gamma parameter controls the width of the RBF kernel. A lower gamma value results in a smoother '
                  'decision boundary, while a higher gamma value makes the decision boundary more complex.')

# Insert SVM accuracy chart
doc.add_heading('SVM Accuracy (All Values)', level=2)
doc.add_paragraph('The chart below shows the accuracy for all gamma values in SVM.')
doc.add_picture(svm_all_chart, width=Inches(5.0))

# Insert SVM accuracy chart excluding gamma=100
doc.add_heading('SVM Accuracy (Excluding Gamma=100)', level=2)
doc.add_paragraph('The chart below shows the accuracy for gamma values in SVM, excluding gamma=100.')
doc.add_picture(svm_filtered_chart, width=Inches(5.0))

# Insert SVM line chart
doc.add_heading('SVM Accuracy Line Chart', level=2)
doc.add_paragraph('The line chart below shows the accuracy across different gamma values in SVM.')
doc.add_picture(svm_line_chart, width=Inches(5.0))

# Insert SVM decision boundary figures with comments
for gamma in svm_gamma_values:
    doc.add_heading(f'SVM with RBF Kernel, gamma={gamma}', level=2)
    doc.add_paragraph(f'The decision boundary for SVM with gamma={gamma} is shown below.')
    
    # Path to the figure
    figure_path = os.path.join(project_folder, f'svm_gamma_{gamma}.png')
    doc.add_picture(figure_path, width=Inches(5.0))


# Add a conclusion section
doc.add_heading('Conclusion', level=1)

doc.add_heading('Logistic Regression:', level=2)
doc.add_paragraph('Logistic Regression is a linear model that works well for simple, linearly separable datasets. It’s easy to understand and explain, as it generates straightforward decision boundaries. However, it struggles with non-linear data. The regularization parameter (C) controls model complexity: small values simplify the model, while large values increase complexity, possibly leading to overfitting. It is a great option when interpretability is important, especially for business or healthcare applications where understanding feature importance is key.')

doc.add_heading('Support Vector Machines (SVM):', level=2)
doc.add_paragraph('SVM with the RBF kernel excels at capturing non-linear patterns in data, making it powerful for complex datasets. The gamma parameter controls how flexible the decision boundary is: low gamma produces smooth boundaries, while high gamma can result in overfitting by creating highly detailed boundaries around data points. SVM is highly effective for datasets where non-linear relationships are crucial, but it can be computationally expensive and harder to explain to non-technical stakeholders due to its complexity.')

doc.add_heading('Recommendation for Real-World Use:', level=2)
doc.add_paragraph('For most real-world use cases, Logistic Regression is preferred when simplicity, speed, and interpretability are important, especially for explaining results to non-technical audiences. However, for complex problems with non-linear data, SVM with RBF kernel offers more flexibility, though it may be harder to explain and requires careful tuning to avoid overfitting.')

# Save the document in the "Small Project 1" folder
report_path = os.path.join(project_folder, 'Predictor_Comparison_Report.docx')
doc.save(report_path)

print("Report generated and saved successfully.")